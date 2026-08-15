import json
import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from cvbankas_tracker.io_utils import ProfileFileReader
from cvbankas_tracker.profile_builder import (
    ProfileGenerationError,
    build_profile_prompt,
    extract_cv_text,
    generate_profile_dict,
    normalize_profile_data,
    write_profile_json,
)

SAMPLE_CV = (
    "Jane Doe\nSenior Python Developer\n"
    "8 years building FastAPI backends, SQL pipelines, Docker, AWS.\n"
    "Open to remote roles.\n"
)

MODEL_OUTPUT = {
    "name": "Jane Doe",
    "target_roles": ["Python Developer", "Backend Developer"],
    "skills": ["python", "fastapi", "sql", "docker", "aws", "python"],
    "preferred_locations": ["Remote"],
    "experience_level": "Senior",
    "years_of_experience": 8,
    "salary_expectation": None,
    "must_have_skills": ["python", "fastapi"],
    "nice_to_have_skills": ["aws"],
    "excluded_keywords": ["warehouse"],
    "additional_keywords": ["backend"],
}


class ExtractCVTextTests(unittest.TestCase):
    def test_reads_txt(self) -> None:
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "cv.txt"
            path.write_text(SAMPLE_CV, encoding="utf-8")
            self.assertIn("FastAPI", extract_cv_text(path))

    def test_reads_md(self) -> None:
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "cv.md"
            path.write_text("# CV\nPython developer", encoding="utf-8")
            self.assertIn("Python developer", extract_cv_text(path))

    def test_missing_file_raises(self) -> None:
        with self.assertRaises(ProfileGenerationError):
            extract_cv_text("does_not_exist.txt")

    def test_unsupported_suffix_raises(self) -> None:
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "cv.rtf"
            path.write_text("x", encoding="utf-8")
            with self.assertRaises(ProfileGenerationError):
                extract_cv_text(path)

    def test_empty_file_raises(self) -> None:
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "cv.txt"
            path.write_text("   \n", encoding="utf-8")
            with self.assertRaises(ProfileGenerationError):
                extract_cv_text(path)

    def test_reads_docx(self) -> None:
        import docx

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "cv.docx"
            document = docx.Document()
            document.add_paragraph("Jane Doe")
            document.add_paragraph("Senior Python Developer with FastAPI")
            document.save(str(path))
            text = extract_cv_text(path)
            self.assertIn("FastAPI", text)

    def test_reads_pdf(self) -> None:
        from pypdf import PdfWriter

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "cv.pdf"
            writer = PdfWriter()
            writer.add_blank_page(width=200, height=200)
            with path.open("wb") as handle:
                writer.write(handle)
            # A blank PDF has no text -> should raise the empty-text error.
            with self.assertRaises(ProfileGenerationError):
                extract_cv_text(path)


class NormalizeProfileDataTests(unittest.TestCase):
    def test_full_payload_dedupes_and_types(self) -> None:
        result = normalize_profile_data(MODEL_OUTPUT)
        self.assertEqual(result["name"], "Jane Doe")
        self.assertEqual(result["skills"].count("python"), 1)
        self.assertEqual(result["years_of_experience"], 8)
        self.assertIsNone(result["salary_expectation"])

    def test_missing_keys_get_defaults(self) -> None:
        result = normalize_profile_data({})
        self.assertEqual(result["name"], "Candidate")
        self.assertEqual(result["experience_level"], "Mid")
        self.assertEqual(result["target_roles"], [])
        self.assertIsNone(result["years_of_experience"])

    def test_invalid_years_becomes_none(self) -> None:
        self.assertIsNone(normalize_profile_data({"years_of_experience": "lots"})["years_of_experience"])

    def test_normalized_profile_loads_via_reader(self) -> None:
        profile = normalize_profile_data(MODEL_OUTPUT)
        with TemporaryDirectory() as tmp:
            path = write_profile_json(Path(tmp) / "p.json", profile)
            loaded = ProfileFileReader().read(path)
        self.assertEqual(loaded.name, "Jane Doe")
        self.assertIn("python", loaded.skills)


class GenerateProfileDictTests(unittest.TestCase):
    def test_rejects_non_ai_backend(self) -> None:
        with self.assertRaises(ProfileGenerationError):
            generate_profile_dict(SAMPLE_CV, backend="rule")
        with self.assertRaises(ProfileGenerationError):
            generate_profile_dict(SAMPLE_CV, backend="demo")

    def test_claude_backend_parses(self) -> None:
        with patch(
            "cvbankas_tracker.profile_builder.run_claude_cli",
            return_value=json.dumps(MODEL_OUTPUT),
        ) as run:
            result = generate_profile_dict(SAMPLE_CV, backend="claude_cli")
        self.assertEqual(result["name"], "Jane Doe")
        run.assert_called_once()

    def test_codex_backend_parses(self) -> None:
        with patch(
            "cvbankas_tracker.profile_builder.run_codex_cli",
            return_value="```json\n" + json.dumps(MODEL_OUTPUT) + "\n```",
        ):
            result = generate_profile_dict(SAMPLE_CV, backend="codex_cli")
        self.assertEqual(result["experience_level"], "Senior")

    def test_prompt_includes_cv_and_schema(self) -> None:
        prompt = build_profile_prompt(SAMPLE_CV)
        self.assertIn("FastAPI", prompt)
        self.assertIn("target_roles", prompt)


if __name__ == "__main__":
    unittest.main()
